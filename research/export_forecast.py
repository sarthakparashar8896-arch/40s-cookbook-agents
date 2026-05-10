"""
Export SKU Demand Forecast to Excel (.xlsx) and Word (.docx)
The 40's Cookbook — Year 1 Planning Model (Jan–Dec 2027)
"""

import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, numbers
)
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, Reference
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────
# CORE DATA
# ─────────────────────────────────────────────

MONTHS = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
          "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

MONTHLY_TOTAL = [5000, 5500, 6000, 7000, 8000, 8500,
                 9000, 9500, 9000, 11000, 12000, 10000]

SKUS = [
    ("Mango Classic",         0.20, "Mango"),
    ("Garlic Ginger",         0.13, "Non-Seasonal"),
    ("Mango Chilli",          0.12, "Mango"),
    ("Lemon Sweet & Sour",    0.10, "Non-Seasonal"),
    ("Stuffed Green Chillies",0.09, "Non-Seasonal"),
    ("Mango Sweet & Sour",    0.09, "Mango"),
    ("Lemon Green Chillies",  0.08, "Non-Seasonal"),
    ("Mustard Chillies",      0.07, "Non-Seasonal"),
    ("Mango Hing",            0.07, "Mango"),
    ("Karonda",               0.05, "Karonda"),
]

SIZE_SPLIT = {"200g": 0.45, "500g": 0.40, "1kg": 0.15}
AVG_WEIGHT = 0.44  # kg per bottle

# Bottles per SKU per month
def sku_monthly_bottles(weight, month_total):
    return round(weight * month_total)

# Pickle kg per SKU per month
def kg(bottles):
    return round(bottles * AVG_WEIGHT)

# Build monthly bottle table: {sku_name: [12 monthly values]}
monthly_bottles = {}
for name, w, _ in SKUS:
    monthly_bottles[name] = [sku_monthly_bottles(w, m) for m in MONTHLY_TOTAL]

# Build monthly kg table
monthly_kg = {name: [kg(b) for b in monthly_bottles[name]] for name, _, _ in SKUS}

# Annual totals
annual_bottles = {name: sum(monthly_bottles[name]) for name, _, _ in SKUS}
annual_kg      = {name: sum(monthly_kg[name])      for name, _, _ in SKUS}

# Size breakdown (annual)
def size_breakdown(total_bottles):
    return {sz: round(total_bottles * frac) for sz, frac in SIZE_SPLIT.items()}

# Gift box data
GIFT_BOX_MONTHS = {
    "3-Pack":  [35, 70, 35, 35, 35, 35, 35, 53, 53, 245, 420, 263],
    "6-Pack":  [45, 90, 45, 45, 45, 45, 45, 68, 68, 315, 540, 338],
    "10-Pack": [20, 40, 20, 20, 20, 20, 20, 30, 30, 140, 240, 150],
}
GIFT_JAR_WEIGHT_KG = 0.05  # 50g per jar
JARS_PER_PACK = {"3-Pack": 3, "6-Pack": 6, "10-Pack": 10}

gift_jars_monthly = []
for i in range(12):
    jars = sum(GIFT_BOX_MONTHS[pk][i] * JARS_PER_PACK[pk] for pk in GIFT_BOX_MONTHS)
    gift_jars_monthly.append(jars)

gift_boxes_monthly = [sum(GIFT_BOX_MONTHS[pk][i] for pk in GIFT_BOX_MONTHS) for i in range(12)]

# Procurement orders (kg incl. buffer)
PROCUREMENT = {
    # SKU: {quarter/period: kg}
    "Mango Classic":          {"Q2 (Apr–Jun, Annual Run)": 10425},
    "Mango Chilli":           {"Q2 (Apr–Jun, Annual Run)": 6250},
    "Mango Sweet & Sour":     {"Q2 (Apr–Jun, Annual Run)": 4688},
    "Mango Hing":             {"Q2 (Apr–Jun, Annual Run)": 3637},
    "Garlic Ginger":          {"Q1": 1133, "Q2": 1613, "Q3": 1888, "Q4": 2266},
    "Lemon Sweet & Sour":     {"Q1":  871, "Q2": 1242, "Q3": 1453, "Q4": 1742},
    "Stuffed Green Chillies": {"Q1":  785, "Q2": 1117, "Q3": 1307, "Q4": 1568},
    "Lemon Green Chillies":   {"Q1":  697, "Q2":  992, "Q3": 1162, "Q4": 1394},
    "Mustard Chillies":       {"Q1":  610, "Q2":  869, "Q3": 1016, "Q4": 1220},
    "Karonda":                {"Q3 (Jul–Sep, Annual Run)": 2650},
}

ANNUAL_ORDER = {
    "Mango Classic":          10387,
    "Garlic Ginger":           6824,
    "Mango Chilli":            6243,
    "Lemon Sweet & Sour":      5301,
    "Stuffed Green Chillies":  4717,
    "Mango Sweet & Sour":      4717,
    "Lemon Green Chillies":    4110,
    "Mustard Chillies":        3603,
    "Mango Hing":              3601,
    "Karonda":                 2651,
}

# ─────────────────────────────────────────────
# STYLE HELPERS
# ─────────────────────────────────────────────

DARK_OLIVE   = "2D3B2D"   # deep forest green (header bg)
WARM_CREAM   = "FDF6E3"   # page background / alt rows
GOLD_ACCENT  = "C8913A"   # accent / highlight
MANGO_YELLOW = "FFD966"   # mango row highlight
KARONDA_RED  = "F4CCCC"   # karonda row
NON_SEA_BLUE = "CFE2F3"   # non-seasonal row
WHITE        = "FFFFFF"
LIGHT_GREY   = "F5F5F5"

def hdr_fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def cell_fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def thin_border():
    s = Side(border_style="thin", color="BBBBBB")
    return Border(left=s, right=s, top=s, bottom=s)

def thick_bottom():
    s = Side(border_style="medium", color="888888")
    return Border(bottom=s)

def hdr_font(size=10, bold=True, color=WHITE):
    return Font(name="Calibri", size=size, bold=bold, color=color)

def body_font(size=9, bold=False, color="222222"):
    return Font(name="Calibri", size=size, bold=bold, color=color)

def style_header_row(ws, row, col_start, col_end, bg=DARK_OLIVE, fg=WHITE, size=10):
    for col in range(col_start, col_end + 1):
        c = ws.cell(row=row, column=col)
        c.fill  = hdr_fill(bg)
        c.font  = hdr_font(size=size, color=fg)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = thin_border()

def style_body_cell(ws, row, col, bg=WHITE, bold=False, align="center"):
    c = ws.cell(row=row, column=col)
    c.fill      = cell_fill(bg)
    c.font      = body_font(bold=bold)
    c.alignment = Alignment(horizontal=align, vertical="center")
    c.border    = thin_border()

def sku_row_color(category):
    return {
        "Mango": MANGO_YELLOW,
        "Karonda": KARONDA_RED,
        "Non-Seasonal": LIGHT_GREY,
    }.get(category, WHITE)

def write_title(ws, text, row, col_start, col_end, size=13):
    ws.merge_cells(start_row=row, start_column=col_start,
                   end_row=row,   end_column=col_end)
    c = ws.cell(row=row, column=col_start, value=text)
    c.font      = Font(name="Calibri", size=size, bold=True, color=DARK_OLIVE)
    c.alignment = Alignment(horizontal="left", vertical="center")

def write_subtitle(ws, text, row, col_start, col_end):
    ws.merge_cells(start_row=row, start_column=col_start,
                   end_row=row,   end_column=col_end)
    c = ws.cell(row=row, column=col_start, value=text)
    c.font      = Font(name="Calibri", size=9, italic=True, color="666666")
    c.alignment = Alignment(horizontal="left")

def set_col_widths(ws, widths):
    for col_idx, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(col_idx)].width = w

def freeze(ws, cell="B3"):
    ws.freeze_panes = cell

# ─────────────────────────────────────────────
# EXCEL EXPORT
# ─────────────────────────────────────────────

def build_excel():
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    _sheet_summary(wb)
    _sheet_monthly_bottles(wb)
    _sheet_monthly_kg(wb)
    _sheet_size_breakdown(wb)
    _sheet_gift_boxes(wb)
    _sheet_mango_strategy(wb)
    _sheet_procurement_calendar(wb)
    _sheet_quarterly_summary(wb)

    path = os.path.join(OUT_DIR, "40s-Cookbook-SKU-Demand-Forecast-v1.xlsx")
    wb.save(path)
    print(f"Excel saved → {path}")
    return path


def _sheet_summary(wb):
    ws = wb.create_sheet("📊 Summary")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 8

    write_title(ws, "The 40's Cookbook — SKU Demand Forecast  |  Year 1: Jan–Dec 2027", 2, 1, 8, size=14)
    write_subtitle(ws, "All volumes are indicative forecasts. Validate against actual sales data after Month 2 and reforecast quarterly.", 3, 1, 8)

    ws.row_dimensions[4].height = 6
    # KPI cards
    kpis = [
        ("Total Bottles (Year 1)", "1,00,500"),
        ("Total Pickle Weight",    "44,221 kg"),
        ("Total Gift Boxes",       "3,750 units"),
        ("Peak Month",             "Nov — 12,000 btl"),
        ("Mango Procurement Run",  "25,000 kg (Apr–Jun)"),
        ("Karonda Procurement Run","2,650 kg (Jul–Sep)"),
        ("Break-Even Volume",      "~1,550 btl/month"),
        ("Year 1 Revenue (est.)",  "₹3.5–4.2 Cr"),
    ]
    style_header_row(ws, 5, 1, 8, bg=DARK_OLIVE)
    ws.cell(5, 1).value = "KEY METRICS"
    ws.merge_cells("A5:H5")
    ws.cell(5, 1).alignment = Alignment(horizontal="center", vertical="center")

    for i, (label, value) in enumerate(kpis):
        r = 6 + i
        ws.cell(r, 1).value = label
        ws.cell(r, 1).font  = body_font(bold=True)
        ws.cell(r, 1).fill  = cell_fill(WARM_CREAM)
        ws.cell(r, 1).border = thin_border()
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=4)
        ws.cell(r, 5).value = value
        ws.cell(r, 5).font  = Font(name="Calibri", size=10, bold=True, color=GOLD_ACCENT)
        ws.cell(r, 5).fill  = cell_fill(WARM_CREAM)
        ws.cell(r, 5).border = thin_border()
        ws.merge_cells(start_row=r, start_column=5, end_row=r, end_column=8)

    # SKU table
    ws.row_dimensions[15].height = 6
    r = 16
    headers = ["SKU", "Category", "Weight %", "Annual Bottles",
               "200g (btl)", "500g (btl)", "1kg (btl)", "Annual Pickle (kg)"]
    for ci, h in enumerate(headers, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, 8)
    ws.row_dimensions[r].height = 20

    for name, w, cat in SKUS:
        r += 1
        bg = sku_row_color(cat)
        sb = size_breakdown(annual_bottles[name])
        row_data = [name, cat, f"{int(w*100)}%",
                    annual_bottles[name],
                    sb["200g"], sb["500g"], sb["1kg"],
                    annual_kg[name]]
        for ci, val in enumerate(row_data, 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci <= 2 else "center",
                            bold=(ci == 1))

    # totals row
    r += 1
    totals = ["TOTAL", "", "100%",
              sum(annual_bottles[n] for n, _, _ in SKUS),
              sum(size_breakdown(annual_bottles[n])["200g"] for n, _, _ in SKUS),
              sum(size_breakdown(annual_bottles[n])["500g"] for n, _, _ in SKUS),
              sum(size_breakdown(annual_bottles[n])["1kg"]  for n, _, _ in SKUS),
              sum(annual_kg[n] for n, _, _ in SKUS)]
    for ci, val in enumerate(totals, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(size=10, color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # Legend
    r += 2
    ws.cell(r, 1).value = "LEGEND"
    ws.cell(r, 1).font = hdr_font(color=DARK_OLIVE)
    legend = [
        (MANGO_YELLOW, "Mango SKU — procure once annually Apr–Jun"),
        (KARONDA_RED,  "Karonda — procure once annually Jul–Sep"),
        (LIGHT_GREY,   "Non-Seasonal — quarterly orders"),
    ]
    for i, (color, label) in enumerate(legend):
        ws.cell(r+1+i, 1).fill  = cell_fill(color)
        ws.cell(r+1+i, 1).value = ""
        ws.cell(r+1+i, 1).border = thin_border()
        ws.cell(r+1+i, 2).value = label
        ws.cell(r+1+i, 2).font  = body_font(size=9)
        ws.merge_cells(start_row=r+1+i, start_column=2, end_row=r+1+i, end_column=8)

    set_col_widths(ws, [28, 16, 10, 16, 12, 12, 12, 18])
    freeze(ws, "A6")


def _sheet_monthly_bottles(wb):
    ws = wb.create_sheet("📦 Monthly Bottles")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 6

    write_title(ws, "Monthly Demand Forecast — Bottles by SKU", 2, 1, 15)
    write_subtitle(ws, "All sizes (200g + 500g + 1kg) combined per SKU.", 3, 1, 15)

    r = 5
    headers = ["SKU", "Category"] + MONTHS + ["TOTAL"]
    for ci, h in enumerate(headers, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, len(headers))
    ws.row_dimensions[r].height = 22

    for name, w, cat in SKUS:
        r += 1
        bg = sku_row_color(cat)
        row_vals = [name, cat] + monthly_bottles[name] + [annual_bottles[name]]
        for ci, val in enumerate(row_vals, 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci <= 2 else "center",
                            bold=(ci == 1))

    # Total row
    r += 1
    totals = ["TOTAL", ""] + MONTHLY_TOTAL + [sum(MONTHLY_TOTAL)]
    for ci, val in enumerate(totals, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # Festive annotation
    r += 2
    ws.cell(r, 1).value = "⚑  Festive spike: October (+35% vs trend), November (+50% vs trend) — Diwali season."
    ws.cell(r, 1).font = Font(name="Calibri", size=9, italic=True, color="884400")
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=15)

    # Mango note
    r += 1
    ws.cell(r, 1).value = "⚑  Mango SKUs (rows shaded yellow): January–March volumes require stock pre-produced in Apr–Jun 2026 season, OR launch mango SKUs in April 2027 (recommended)."
    ws.cell(r, 1).font = Font(name="Calibri", size=9, italic=True, color="886600")
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=15)

    set_col_widths(ws, [26, 14] + [7]*12 + [9])
    freeze(ws, "C6")


def _sheet_monthly_kg(wb):
    ws = wb.create_sheet("⚖️ Pickle Weight (kg)")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 6

    write_title(ws, "Monthly Pickle Weight by SKU (kg) — For CM Order Sizing", 2, 1, 15)
    write_subtitle(ws, "Weighted average fill: 0.44 kg/bottle (200g×45% + 500g×40% + 1kg×15%). Use these figures when placing orders with contract manufacturer.", 3, 1, 15)

    r = 5
    headers = ["SKU", "Category"] + MONTHS + ["TOTAL (kg)"]
    for ci, h in enumerate(headers, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, len(headers))
    ws.row_dimensions[r].height = 22

    for name, w, cat in SKUS:
        r += 1
        bg = sku_row_color(cat)
        row_vals = [name, cat] + monthly_kg[name] + [annual_kg[name]]
        for ci, val in enumerate(row_vals, 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci <= 2 else "center",
                            bold=(ci == 1))

    r += 1
    monthly_totals_kg = [round(sum(monthly_kg[n][i] for n, _, _ in SKUS)) for i in range(12)]
    totals = ["TOTAL (kg)", ""] + monthly_totals_kg + [sum(monthly_totals_kg)]
    for ci, val in enumerate(totals, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # Mango subtotal
    r += 1
    mango_names = [n for n, _, c in SKUS if c == "Mango"]
    m_monthly = [round(sum(monthly_kg[n][i] for n in mango_names)) for i in range(12)]
    mango_row = ["Mango SKUs subtotal", "(seasonal)"] + m_monthly + [sum(m_monthly)]
    for ci, val in enumerate(mango_row, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = cell_fill(MANGO_YELLOW)
        ws.cell(r, ci).font  = body_font(bold=True, color="664400")
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 2 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    set_col_widths(ws, [26, 14] + [7]*12 + [10])
    freeze(ws, "C6")


def _sheet_size_breakdown(wb):
    ws = wb.create_sheet("📐 Size Breakdown")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 6

    write_title(ws, "Annual Bottles by SKU × Size", 2, 1, 8)
    write_subtitle(ws, "Size split: 200g = 45%, 500g = 40%, 1kg = 15% (uniform across all SKUs).", 3, 1, 8)

    r = 5
    headers = ["SKU", "Category", "Total Bottles",
               "200g (btl)", "200g (kg)", "500g (btl)", "500g (kg)", "1kg (btl)", "1kg (kg)"]
    for ci, h in enumerate(headers, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, len(headers))
    ws.row_dimensions[r].height = 22

    for name, w, cat in SKUS:
        r += 1
        bg = sku_row_color(cat)
        sb = size_breakdown(annual_bottles[name])
        row_vals = [
            name, cat, annual_bottles[name],
            sb["200g"],  round(sb["200g"] * 0.2),
            sb["500g"],  round(sb["500g"] * 0.5),
            sb["1kg"],   round(sb["1kg"]  * 1.0),
        ]
        for ci, val in enumerate(row_vals, 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci <= 2 else "center",
                            bold=(ci == 1))

    r += 1
    all_sb_200 = sum(size_breakdown(annual_bottles[n])["200g"] for n, _, _ in SKUS)
    all_sb_500 = sum(size_breakdown(annual_bottles[n])["500g"] for n, _, _ in SKUS)
    all_sb_1kg = sum(size_breakdown(annual_bottles[n])["1kg"]  for n, _, _ in SKUS)
    totals = ["TOTAL", "", sum(annual_bottles.values()),
              all_sb_200, round(all_sb_200*0.2),
              all_sb_500, round(all_sb_500*0.5),
              all_sb_1kg, round(all_sb_1kg*1.0)]
    for ci, val in enumerate(totals, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 2 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    set_col_widths(ws, [26, 14, 14, 11, 10, 11, 10, 10, 10])
    freeze(ws, "A6")


def _sheet_gift_boxes(wb):
    ws = wb.create_sheet("🎁 Gift Boxes")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 6

    write_title(ws, "Gift Box Demand Forecast — Monthly", 2, 1, 16)
    write_subtitle(ws, "3-pack (3×50g) | 6-pack (6×50g) | 10-pack (10×50g). Festive spike: Oct–Nov = 50.7% of annual gift box volume.", 3, 1, 16)

    r = 5
    headers = ["Gift Pack", "Jars/Box"] + MONTHS + ["TOTAL Units", "TOTAL Jars", "Pickle (kg)"]
    for ci, h in enumerate(headers, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, len(headers))
    ws.row_dimensions[r].height = 22

    pack_colors = {"3-Pack": "EAF4E1", "6-Pack": "D5EBF5", "10-Pack": "F5E6FF"}
    for pk, jars in [("3-Pack", 3), ("6-Pack", 6), ("10-Pack", 10)]:
        r += 1
        total_units = sum(GIFT_BOX_MONTHS[pk])
        total_jars  = total_units * jars
        total_kg    = round(total_jars * GIFT_JAR_WEIGHT_KG, 1)
        row_vals    = [pk, jars] + GIFT_BOX_MONTHS[pk] + [total_units, total_jars, total_kg]
        bg = pack_colors[pk]
        for ci, val in enumerate(row_vals, 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci == 1 else "center",
                            bold=(ci == 1))

    r += 1
    totals_units = [sum(GIFT_BOX_MONTHS[pk][i] for pk in GIFT_BOX_MONTHS) for i in range(12)]
    totals_jars  = [sum(GIFT_BOX_MONTHS[pk][i] * JARS_PER_PACK[pk] for pk in GIFT_BOX_MONTHS) for i in range(12)]
    t_units = sum(totals_units)
    t_jars  = sum(totals_jars)
    row_vals = ["TOTAL", ""] + totals_units + [t_units, t_jars, round(t_jars*GIFT_JAR_WEIGHT_KG, 1)]
    for ci, val in enumerate(row_vals, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 1 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # 50g jar SKU breakdown
    r += 3
    write_title(ws, "50g Jar SKU Breakdown (Annual, across all gift packs)", r, 1, 6, size=11)
    r += 1
    hdrs = ["SKU", "In 3-Pack", "In 6-Pack", "In 10-Pack", "Total 50g Jars", "Pickle (kg)"]
    for ci, h in enumerate(hdrs, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, 6)

    jar_data = [
        ("Mango Classic",         1314, 1689, 749, 3752, 188),
        ("Garlic Ginger",         1314, 1689, 749, 3752, 188),
        ("Mango Chilli",             0, 1689, 749, 2438, 122),
        ("Lemon Sweet & Sour",    1314, 1689, 749, 3752, 188),
        ("Stuffed Green Chillies",   0, 1689, 749, 2438, 122),
        ("Mango Sweet & Sour",       0, 1689, 749, 2438, 122),
        ("Lemon Green Chillies",     0,    0, 749,  749,  37),
        ("Mustard Chillies",         0,    0, 749,  749,  37),
        ("Mango Hing",               0,    0, 749,  749,  37),
        ("Karonda",                  0,    0, 749,  749,  37),
    ]
    for row_d in jar_data:
        r += 1
        name = row_d[0]
        cat  = next(c for n, w, c in SKUS if n == name)
        bg   = sku_row_color(cat)
        for ci, val in enumerate(row_d, 1):
            ws.cell(r, ci).value = val if val != 0 else "—"
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci == 1 else "center",
                            bold=(ci == 1))

    r += 1
    totals_jar = ["TOTAL", 3942, 10134, 7490, 21566, 1078]
    for ci, val in enumerate(totals_jar, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 1 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    set_col_widths(ws, [20, 8] + [7]*12 + [12, 10, 10])
    freeze(ws, "B6")


def _sheet_mango_strategy(wb):
    ws = wb.create_sheet("🥭 Mango Procurement")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 6

    write_title(ws, "Mango Seasonality — Annual Procurement Strategy", 2, 1, 7, size=13)
    write_subtitle(ws, "Raw mango season: April–June. All mango pickle for the year must be contracted and produced in this window.", 3, 1, 7)

    r = 5
    write_title(ws, "Annual Mango Requirement", r, 1, 7, size=11)
    r += 1
    hdrs = ["SKU", "Annual Bottles", "Annual Base (kg)", "Gift Jars (kg)", "Buffer 15% (kg)", "TOTAL ORDER (kg)", "% of Mango Run"]
    for ci, h in enumerate(hdrs, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, 7)

    mango_skus = [(n, w, c) for n, w, c in SKUS if c == "Mango"]
    mango_gift = {"Mango Classic": 188, "Mango Chilli": 122, "Mango Sweet & Sour": 122, "Mango Hing": 37}
    mango_order = {"Mango Classic": 10425, "Mango Chilli": 6250, "Mango Sweet & Sour": 4688, "Mango Hing": 3637}
    total_mango_order = sum(mango_order.values())

    for name, w, cat in mango_skus:
        r += 1
        base = annual_kg[name]
        gift = mango_gift[name]
        order = mango_order[name]
        buf  = order - base - gift
        pct  = f"{round(order/total_mango_order*100, 1)}%"
        for ci, val in enumerate([name, annual_bottles[name], base, gift, buf, order, pct], 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=MANGO_YELLOW,
                            align="left" if ci == 1 else "center", bold=(ci == 1))

    r += 1
    totals = ["TOTAL", 48240, 21226, 469, sum(mango_order[n]-annual_kg[n]-mango_gift[n] for n, _, _ in mango_skus),
              total_mango_order, "100%"]
    for ci, val in enumerate(totals, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 1 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # Production schedule
    r += 3
    write_title(ws, "Production Run Schedule (Apr–Jun 2027)", r, 1, 5, size=11)
    r += 1
    for ci, h in enumerate(["Month", "Production (kg)", "Cumulative (kg)", "% of Run", "Notes"], 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, 5)

    prod_schedule = [
        ("April 2027",  10000, 10000, "40%", "Peak raw mango; start production"),
        ("May 2027",    10000, 20000, "40%", "Peak season; highest quality"),
        ("June 2027",    5000, 25000, "20%", "End of season; finalise run"),
    ]
    for name, qty, cum, pct, note in prod_schedule:
        r += 1
        for ci, val in enumerate([name, qty, cum, pct, note], 1):
            style_body_cell(ws, r, ci, bg=MANGO_YELLOW,
                            align="left" if ci in [1, 5] else "center")
            ws.cell(r, ci).value = val

    r += 1
    for ci, val in enumerate(["TOTAL", 25000, "", "100%", ""], 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 1 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # Drawdown table
    r += 3
    write_title(ws, "Monthly Stock Drawdown from 25,000 kg Season Stock", r, 1, 5, size=11)
    r += 1
    for ci, h in enumerate(["Month", "Demand (kg)", "Closing Stock (kg)", "Buffer Status", "Action"], 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, 5)

    mango_monthly_demand = [sum(monthly_kg[n][i] for n, _, c in SKUS if c == "Mango") for i in range(12)]
    stock = 25000
    months_ext = ["Apr 2027", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec",
                  "Jan 2028*", "Feb 2028*", "Mar 2028*"]
    demands_ext = mango_monthly_demand[3:] + [1200, 1350, 1450]

    for i, (mo, dem) in enumerate(zip(months_ext, demands_ext)):
        stock -= dem
        r += 1
        status = "✓ Healthy" if stock > 5000 else ("⚠ Low" if stock > 2000 else "🔴 Critical")
        action = "" if stock > 8000 else ("Monitor closely" if stock > 4000 else "Place Year 2 order NOW")
        bg = LIGHT_GREY if i % 2 == 0 else WHITE
        for ci, val in enumerate([mo, dem, stock, status, action], 1):
            style_body_cell(ws, r, ci, bg=bg, align="left" if ci in [1, 4, 5] else "center")
            ws.cell(r, ci).value = val

    r += 2
    ws.cell(r, 1).value = "* Jan–Mar 2028 are estimates. Place April 2028 mango order by mid-March 2028."
    ws.cell(r, 1).font = Font(name="Calibri", size=9, italic=True, color="666666")
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)

    set_col_widths(ws, [18, 18, 17, 18, 12, 14, 22])
    freeze(ws, "A6")


def _sheet_procurement_calendar(wb):
    ws = wb.create_sheet("📅 Procurement Calendar")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 6

    write_title(ws, "Master Procurement Calendar — Monthly Orders (kg, incl. buffer)", 2, 1, 13)
    write_subtitle(ws, "Order quantities include 15–20% safety buffer. Place orders 3–4 weeks before delivery month. Mango and Karonda are annual runs; all others quarterly.", 3, 1, 13)

    r = 5
    col_headers = ["Action / SKU", "Type"] + MONTHS + ["YEAR TOTAL"]
    for ci, h in enumerate(col_headers, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, len(col_headers))
    ws.row_dimensions[r].height = 22

    cal_rows = []

    # Non-seasonal quarterly
    q_data = {
        "Garlic Ginger":          [1133, 0, 0, 1613, 0, 0, 1888, 0, 0, 2266, 0, 0],
        "Lemon Sweet & Sour":     [ 871, 0, 0, 1242, 0, 0, 1453, 0, 0, 1742, 0, 0],
        "Stuffed Green Chillies": [ 785, 0, 0, 1117, 0, 0, 1307, 0, 0, 1568, 0, 0],
        "Lemon Green Chillies":   [ 697, 0, 0,  992, 0, 0, 1162, 0, 0, 1394, 0, 0],
        "Mustard Chillies":       [ 610, 0, 0,  869, 0, 0, 1016, 0, 0, 1220, 0, 0],
    }
    # Mango (Apr–Jun split 40/40/20 of 25000)
    mango_monthly_orders = {
        "Mango Classic":      [0,0,0, 4170,4170,2085, 0,0,0, 0,0,0],
        "Mango Chilli":       [0,0,0, 2500,2500,1250, 0,0,0, 0,0,0],
        "Mango Sweet & Sour": [0,0,0, 1875,1875, 938, 0,0,0, 0,0,0],
        "Mango Hing":         [0,0,0, 1455,1455, 727, 0,0,0, 0,0,0],
    }
    # Karonda (Jul–Sep)
    karonda_monthly = [0,0,0, 0,0,0, 1060,1060,530, 0,0,0]

    all_order_rows = {}
    for n in [x[0] for x in SKUS]:
        if n in q_data:
            all_order_rows[n] = q_data[n]
        elif n in mango_monthly_orders:
            all_order_rows[n] = mango_monthly_orders[n]
        else:
            all_order_rows[n] = karonda_monthly

    for name, w, cat in SKUS:
        r += 1
        bg = sku_row_color(cat)
        vals = all_order_rows[name]
        display = [v if v > 0 else "" for v in vals]
        total = sum(vals)
        row_vals = [name, cat] + display + [total]
        for ci, val in enumerate(row_vals, 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci <= 2 else "center",
                            bold=(ci == 1 or ci == len(row_vals)))

    # Monthly totals
    r += 1
    monthly_col_totals = []
    for mi in range(12):
        col_total = sum(all_order_rows[n][mi] for n, _, _ in SKUS)
        monthly_col_totals.append(col_total if col_total > 0 else "")
    year_total = sum(v for v in monthly_col_totals if isinstance(v, int))
    totals_row = ["TOTAL INBOUND (kg)", ""] + [v for v in monthly_col_totals] + [year_total]
    for ci, val in enumerate(totals_row, 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 2 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # Milestone row
    r += 2
    milestones = {
        1:  "Q1 order received",
        4:  "🥭 Mango delivery starts",
        7:  "🔴 Karonda season",
        10: "🎆 Diwali — peak demand",
        11: "Peak month — ensure full stock",
    }
    ws.cell(r, 1).value = "Key Events"
    ws.cell(r, 1).font  = hdr_font(color=DARK_OLIVE, size=9)
    for mi, label in milestones.items():
        ws.cell(r, mi+2).value = label
        ws.cell(r, mi+2).font  = Font(name="Calibri", size=8, italic=True, color=GOLD_ACCENT)
        ws.cell(r, mi+2).alignment = Alignment(horizontal="center", wrap_text=True)

    set_col_widths(ws, [26, 14] + [7]*12 + [11])
    freeze(ws, "C6")


def _sheet_quarterly_summary(wb):
    ws = wb.create_sheet("📋 Quarterly Summary")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 6

    write_title(ws, "Quarterly & Annual Summary — Procurement Orders (kg)", 2, 1, 8)
    write_subtitle(ws, "All figures include safety buffer (15–20%). Use this as the master order reference for planning conversations with your CM.", 3, 1, 8)

    r = 5
    headers = ["SKU", "Category", "Q1\nJan–Mar", "Q2\nApr–Jun", "Q3\nJul–Sep", "Q4\nOct–Dec", "YEAR TOTAL (kg)", "Order Type", "Order Placement"]
    for ci, h in enumerate(headers, 1):
        ws.cell(r, ci).value = h
    style_header_row(ws, r, 1, len(headers))
    ws.row_dimensions[r].height = 30

    sku_qtly = {
        "Mango Classic":          (0,     10425, 0,    0,    "Annual (Apr–Jun)", "Mid-March"),
        "Garlic Ginger":          (1133,   1613, 1888, 2266, "Quarterly",        "Dec/Mar/Jun/Sep"),
        "Mango Chilli":           (0,      6250, 0,    0,    "Annual (Apr–Jun)", "Mid-March"),
        "Lemon Sweet & Sour":     (871,    1242, 1453, 1742, "Quarterly",        "Dec/Mar/Jun/Sep"),
        "Stuffed Green Chillies": (785,    1117, 1307, 1568, "Quarterly",        "Dec/Mar/Jun/Sep"),
        "Mango Sweet & Sour":     (0,      4688, 0,    0,    "Annual (Apr–Jun)", "Mid-March"),
        "Lemon Green Chillies":   (697,     992, 1162, 1394, "Quarterly",        "Dec/Mar/Jun/Sep"),
        "Mustard Chillies":       (610,     869, 1016, 1220, "Quarterly",        "Dec/Mar/Jun/Sep"),
        "Mango Hing":             (0,      3637, 0,    0,    "Annual (Apr–Jun)", "Mid-March"),
        "Karonda":                (0,         0, 2650, 0,    "Annual (Jul–Sep)", "Late-June"),
    }

    for name, w, cat in SKUS:
        r += 1
        bg = sku_row_color(cat)
        q1, q2, q3, q4, otype, oplacement = sku_qtly[name]
        total = q1 + q2 + q3 + q4
        for ci, val in enumerate([name, cat,
                                   q1 or "—", q2 or "—", q3 or "—", q4 or "—",
                                   total, otype, oplacement], 1):
            ws.cell(r, ci).value = val
            style_body_cell(ws, r, ci, bg=bg,
                            align="left" if ci in [1, 2, 8, 9] else "center",
                            bold=(ci == 1 or ci == 7))

    r += 1
    q1t = sum(sku_qtly[n][0] for n, _, _ in SKUS)
    q2t = sum(sku_qtly[n][1] for n, _, _ in SKUS)
    q3t = sum(sku_qtly[n][2] for n, _, _ in SKUS)
    q4t = sum(sku_qtly[n][3] for n, _, _ in SKUS)
    for ci, val in enumerate(["TOTAL (kg)", "", q1t, q2t, q3t, q4t, q1t+q2t+q3t+q4t, "", ""], 1):
        ws.cell(r, ci).value = val
        ws.cell(r, ci).fill  = hdr_fill(DARK_OLIVE)
        ws.cell(r, ci).font  = hdr_font(color=WHITE)
        ws.cell(r, ci).alignment = Alignment(horizontal="center" if ci > 2 else "left", vertical="center")
        ws.cell(r, ci).border = thin_border()

    # Notes
    r += 2
    notes = [
        "Q1 (Jan–Mar): Non-seasonal SKUs only. Mango SKUs launch April (recommended).",
        "Q2 (Apr–Jun): Largest quarter by procurement — full annual mango run (25,000 kg) delivered here.",
        "Q3 (Jul–Sep): Karonda annual run (2,650 kg). Full gift box inventory must arrive by Sep 30.",
        "Q4 (Oct–Dec): Peak sales quarter (Diwali). No new seasonal runs — ship from warehouse stock.",
    ]
    for note in notes:
        ws.cell(r, 1).value = note
        ws.cell(r, 1).font  = Font(name="Calibri", size=9, italic=True, color="444444")
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=9)
        r += 1

    set_col_widths(ws, [26, 14, 10, 10, 10, 10, 14, 18, 20])
    freeze(ws, "A6")


# ─────────────────────────────────────────────
# WORD EXPORT
# ─────────────────────────────────────────────

def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"),   "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"),  hex_color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def doc_add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*hex_to_rgb(DARK_OLIVE))
    return p

def doc_add_para(doc, text, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p

def doc_table(doc, headers, rows, col_widths=None, row_colors=None):
    """Build a styled table."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    hdr_row.height = Cm(0.7)
    for ci, h in enumerate(headers):
        cell = hdr_row.cells[ci]
        set_cell_bg(cell, DARK_OLIVE)
        set_cell_borders(cell)
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(str(h))
        run.bold = True
        run.font.size  = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        bg = row_colors[ri] if row_colors else ("F9F9F9" if ri % 2 == 0 else WHITE)
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.bold = (ci == 0)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def sku_doc_color(cat):
    return {
        "Mango": "FFF2CC",
        "Karonda": "F4CCCC",
        "Non-Seasonal": "F5F5F5",
    }.get(cat, WHITE)

def build_word():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Title
    title = doc.add_heading("The 40's Cookbook", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(*hex_to_rgb(DARK_OLIVE))
        run.font.size = Pt(22)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph()
    run = sub.add_run("SKU Demand Forecast & Procurement Calendar  |  Year 1: Jan–Dec 2027")
    run.font.size  = Pt(12)
    run.font.color.rgb = RGBColor(*hex_to_rgb(GOLD_ACCENT))
    run.bold = True

    doc.add_paragraph()
    doc_add_para(doc, "All volumes are indicative forecasts. Validate against actual sales data after Month 2 and reforecast quarterly.", italic=True, size=9)
    doc.add_paragraph()

    # ── Section 1: TL;DR ──
    doc_add_heading(doc, "TL;DR — Key Metrics", 1)
    kpis = [
        ("Total Bottles (Year 1)",       "1,00,500"),
        ("Total Pickle Weight",          "44,221 kg"),
        ("Total Gift Boxes",             "3,750 units (21,576 mini 50g jars)"),
        ("Peak Month",                   "November (Diwali): 12,000 bottles"),
        ("Mango Procurement Run",        "25,000 kg  —  April–June 2027"),
        ("Karonda Procurement Run",      "2,650 kg  —  July–September 2027"),
        ("Break-Even Volume",            "~1,550 bottles/month"),
        ("Estimated Year 1 Revenue",     "₹3.5–4.2 Cr (blended MRP ₹370–420)"),
    ]
    doc_table(doc,
              ["Metric", "Value"],
              [[k, v] for k, v in kpis],
              col_widths=[8, 10])
    doc.add_paragraph()

    # ── Section 2: Assumptions ──
    doc_add_heading(doc, "Assumptions & Methodology", 1)
    doc_add_para(doc, "Volume trajectory: 5,000 btl/month (Jan) → 12,000 btl/month (Nov Diwali peak) → 10,000 (Dec). 8–10% MoM growth M1–M6, 5–8% thereafter with festive spike in Oct–Nov.")
    doc.add_paragraph()

    sku_rows = [[name, f"{int(w*100)}%", cat] for name, w, cat in SKUS]
    sku_colors = [sku_doc_color(cat) for _, _, cat in SKUS]
    doc_add_heading(doc, "SKU Demand Weights", 2)
    doc_table(doc,
              ["SKU", "Demand Weight", "Category"],
              sku_rows,
              col_widths=[8, 4, 5],
              row_colors=sku_colors)
    doc.add_paragraph()

    doc_add_para(doc, "Size mix (uniform across all SKUs): 200g = 45%  |  500g = 40%  |  1 kg = 15%. Weighted average fill = 0.44 kg/bottle.")
    doc.add_paragraph()

    # ── Section 3: Monthly Volume ──
    doc_add_heading(doc, "Monthly Volume Forecast — Bottles by SKU", 1)
    doc_add_para(doc, "All sizes combined per SKU per month.", italic=True)
    doc.add_paragraph()

    m_headers = ["SKU"] + MONTHS + ["TOTAL"]
    m_rows = []
    m_colors = []
    for name, w, cat in SKUS:
        m_rows.append([name] + monthly_bottles[name] + [annual_bottles[name]])
        m_colors.append(sku_doc_color(cat))
    m_rows.append(["TOTAL"] + MONTHLY_TOTAL + [sum(MONTHLY_TOTAL)])
    m_colors.append(DARK_OLIVE)
    doc_table(doc, m_headers, m_rows,
              col_widths=[5.5] + [1.1]*12 + [1.5],
              row_colors=m_colors)
    doc.add_paragraph()
    doc_add_para(doc, "Yellow rows = Mango SKUs (seasonal). Pink row = Karonda (seasonal). Grey = Non-Seasonal.", italic=True, size=8)
    doc.add_paragraph()

    # ── Section 4: Pickle Weight ──
    doc_add_heading(doc, "Monthly Pickle Weight by SKU (kg)", 1)
    doc_add_para(doc, "Use these figures for CM order sizing. Based on weighted average fill of 0.44 kg/bottle.", italic=True)
    doc.add_paragraph()

    kg_headers = ["SKU"] + MONTHS + ["TOTAL kg"]
    kg_rows = []
    kg_colors = []
    for name, w, cat in SKUS:
        kg_rows.append([name] + monthly_kg[name] + [annual_kg[name]])
        kg_colors.append(sku_doc_color(cat))
    monthly_kg_totals = [round(sum(monthly_kg[n][i] for n, _, _ in SKUS)) for i in range(12)]
    kg_rows.append(["TOTAL"] + monthly_kg_totals + [sum(monthly_kg_totals)])
    kg_colors.append(DARK_OLIVE)
    doc_table(doc, kg_headers, kg_rows,
              col_widths=[5.5] + [1.1]*12 + [1.5],
              row_colors=kg_colors)
    doc.add_paragraph()

    # ── Section 5: Size Breakdown ──
    doc_add_heading(doc, "Annual Bottles by SKU × Size", 1)
    sb_headers = ["SKU", "Total Bottles", "200g (btl)", "200g (kg)", "500g (btl)", "500g (kg)", "1kg (btl)", "1kg (kg)"]
    sb_rows = []
    sb_colors = []
    for name, w, cat in SKUS:
        sb = size_breakdown(annual_bottles[name])
        sb_rows.append([name, annual_bottles[name],
                        sb["200g"], round(sb["200g"]*0.2),
                        sb["500g"], round(sb["500g"]*0.5),
                        sb["1kg"],  round(sb["1kg"]*1.0)])
        sb_colors.append(sku_doc_color(cat))
    doc_table(doc, sb_headers, sb_rows,
              col_widths=[6, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5],
              row_colors=sb_colors)
    doc.add_paragraph()

    # ── Section 6: Gift Boxes ──
    doc_add_heading(doc, "Gift Box Demand Forecast", 1)
    doc_add_para(doc, "3-Pack (3×50g)  |  6-Pack (6×50g)  |  10-Pack (10×50g). October–November accounts for 50.7% of annual gift volume (Diwali season). All gift box stock must be warehoused by September 30.")
    doc.add_paragraph()

    gb_headers = ["Gift Pack", "Jars/Box"] + MONTHS + ["Total Units", "Total Jars", "Pickle (kg)"]
    gb_rows = []
    gb_colors = []
    pack_doc_colors = {"3-Pack": "EAF4E1", "6-Pack": "D5EBF5", "10-Pack": "F5E6FF"}
    for pk, jars in [("3-Pack", 3), ("6-Pack", 6), ("10-Pack", 10)]:
        tu = sum(GIFT_BOX_MONTHS[pk])
        tj = tu * jars
        row = [pk, jars] + GIFT_BOX_MONTHS[pk] + [tu, tj, round(tj*0.05, 1)]
        gb_rows.append(row)
        gb_colors.append(pack_doc_colors[pk])
    tu_all = sum(sum(GIFT_BOX_MONTHS[pk]) for pk in GIFT_BOX_MONTHS)
    _totals_units = [sum(GIFT_BOX_MONTHS[pk][i] for pk in GIFT_BOX_MONTHS) for i in range(12)]
    _totals_jars  = [sum(GIFT_BOX_MONTHS[pk][i] * JARS_PER_PACK[pk] for pk in GIFT_BOX_MONTHS) for i in range(12)]
    tj_all = sum(_totals_jars)
    gb_rows.append(["TOTAL", ""] + [str(v) for v in _totals_units] + [tu_all, tj_all, round(tj_all*0.05, 1)])
    gb_colors.append(DARK_OLIVE)
    doc_table(doc, gb_headers, gb_rows,
              col_widths=[2, 1.5] + [1.0]*12 + [2, 2, 2],
              row_colors=gb_colors)
    doc.add_paragraph()

    # ── Section 7: Mango Strategy ──
    doc_add_heading(doc, "Mango Seasonality — Procurement Strategy", 1)
    doc_add_para(doc, "Raw green mango is only available April–June. All annual mango pickle must be contracted and produced in a single seasonal run. The brand must place the CM order by mid-March.")
    doc.add_paragraph()

    doc_add_heading(doc, "Recommendation: Launch mango SKUs in April 2027", 2)
    doc_add_para(doc, "Launch January 2027 with 6 non-mango SKUs. Add all 4 mango SKUs in April as a 'Mango Season Launch' campaign event — this avoids the need for pre-season stock and turns the seasonal constraint into a brand moment.")
    doc.add_paragraph()

    mango_headers = ["SKU", "Annual Bottles", "Base (kg)", "Gift (kg)", "Buffer (kg)", "ORDER (kg)", "% of Run"]
    mango_rows = []
    mango_skus_list = [(n, w, c) for n, w, c in SKUS if c == "Mango"]
    mango_gift_d = {"Mango Classic": 188, "Mango Chilli": 122, "Mango Sweet & Sour": 122, "Mango Hing": 37}
    mango_order_d = {"Mango Classic": 10425, "Mango Chilli": 6250, "Mango Sweet & Sour": 4688, "Mango Hing": 3637}
    total_mo = sum(mango_order_d.values())
    for name, w, cat in mango_skus_list:
        base  = annual_kg[name]
        gift  = mango_gift_d[name]
        order = mango_order_d[name]
        buf   = order - base - gift
        mango_rows.append([name, annual_bottles[name], base, gift, buf, order, f"{round(order/total_mo*100,1)}%"])
    mango_rows.append(["TOTAL", 48240, 21226, 469, sum(mango_order_d[n]-annual_kg[n]-mango_gift_d[n] for n,_,_ in mango_skus_list), total_mo, "100%"])
    doc_table(doc, mango_headers, mango_rows,
              col_widths=[5.5, 2.5, 2.5, 2, 2.5, 2.5, 2],
              row_colors=[sku_doc_color("Mango")]*len(mango_rows))
    doc.add_paragraph()

    doc_add_heading(doc, "Production Run: April–June 2027  (Total: 25,000 kg)", 2)
    prod_rows = [
        ["April 2027",  "10,000 kg", "40%", "Peak availability; start first batch"],
        ["May 2027",    "10,000 kg", "40%", "Peak quality; main run"],
        ["June 2027",    "5,000 kg", "20%", "End of season; finalise"],
        ["TOTAL",       "25,000 kg", "100%", "Covers Apr 2027 – Mar 2028 + buffer"],
    ]
    doc_table(doc, ["Month", "Qty (kg)", "% of Run", "Notes"],
              prod_rows, col_widths=[4, 3, 3, 8])
    doc.add_paragraph()

    # ── Section 8: Quarterly Procurement ──
    doc_add_heading(doc, "Quarterly Procurement Orders — All SKUs (kg, incl. buffer)", 1)
    doc.add_paragraph()

    qtly_headers = ["SKU", "Q1 Jan–Mar", "Q2 Apr–Jun", "Q3 Jul–Sep", "Q4 Oct–Dec", "YEAR TOTAL", "Order Type"]
    qtly_data = {
        "Mango Classic":          (0,     10425, 0,    0,    "Annual"),
        "Garlic Ginger":          (1133,   1613, 1888, 2266, "Quarterly"),
        "Mango Chilli":           (0,      6250, 0,    0,    "Annual"),
        "Lemon Sweet & Sour":     (871,    1242, 1453, 1742, "Quarterly"),
        "Stuffed Green Chillies": (785,    1117, 1307, 1568, "Quarterly"),
        "Mango Sweet & Sour":     (0,      4688, 0,    0,    "Annual"),
        "Lemon Green Chillies":   (697,     992, 1162, 1394, "Quarterly"),
        "Mustard Chillies":       (610,     869, 1016, 1220, "Quarterly"),
        "Mango Hing":             (0,      3637, 0,    0,    "Annual"),
        "Karonda":                (0,         0, 2650, 0,    "Annual"),
    }
    qtly_rows = []
    qtly_colors = []
    for name, w, cat in SKUS:
        q1, q2, q3, q4, otype = qtly_data[name]
        qtly_rows.append([name,
                          q1 if q1 else "—",
                          q2 if q2 else "—",
                          q3 if q3 else "—",
                          q4 if q4 else "—",
                          q1+q2+q3+q4, otype])
        qtly_colors.append(sku_doc_color(cat))
    q1t = sum(qtly_data[n][0] for n,_,_ in SKUS)
    q2t = sum(qtly_data[n][1] for n,_,_ in SKUS)
    q3t = sum(qtly_data[n][2] for n,_,_ in SKUS)
    q4t = sum(qtly_data[n][3] for n,_,_ in SKUS)
    qtly_rows.append(["TOTAL (kg)", q1t, q2t, q3t, q4t, q1t+q2t+q3t+q4t, ""])
    qtly_colors.append(DARK_OLIVE)
    doc_table(doc, qtly_headers, qtly_rows,
              col_widths=[6, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5],
              row_colors=qtly_colors)
    doc.add_paragraph()

    # ── Section 9: Milestones ──
    doc_add_heading(doc, "Key Planning Milestones", 1)
    milestone_rows = [
        ["Nov 2026",  "Finalise contract manufacturer for pickle production"],
        ["Dec 2026",  "Place Q1 non-seasonal orders (4,096 kg)"],
        ["Jan 2027",  "LAUNCH — 6 SKUs (non-mango)"],
        ["Mar 2027",  "Place full mango order with CM (25,000 kg) + Q2 non-seasonal orders"],
        ["Apr 2027",  "Mango SKUs go live — 'Mango Season Launch' campaign"],
        ["Jun 2027",  "Place Q3 orders + karonda pre-order"],
        ["Jun 2027",  "Confirm Diwali gift box inventory plan"],
        ["Jul 2027",  "Karonda Season Launch"],
        ["Sep 30",    "ALL Diwali gift box stock must be warehoused"],
        ["Sep 2027",  "Place Q4 orders"],
        ["Dec 2027",  "Place Q1 2028 orders; begin Year 2 mango planning"],
    ]
    doc_table(doc, ["Date", "Action"],
              milestone_rows, col_widths=[3.5, 14])
    doc.add_paragraph()

    # ── Section 10: Open Questions ──
    doc_add_heading(doc, "Open Questions & Validation Required", 1)
    questions = [
        "CM capacity: Can the chosen CM handle a 25,000 kg mango run in Apr–Jun? Confirm by January 2027.",
        "MOQ for 50g gift jars: Indian glass suppliers (AGI, HNG) typically have MOQs of 5,000–10,000 pieces for non-standard sizes.",
        "Shelf life validation: Run accelerated shelf-life tests (6–8 weeks) for each variant before committing to large batch quantities.",
        "Karonda CM: Not all CMs process karonda — confirm sourcing network with your CM partner.",
        "SKU weight assumptions: Validate demand weights after Month 2 sales data; reforecast quarterly.",
        "Stuffed Green Chilli shelf life: Bharwa mirch has faster oil migration; confirm jar sealing and shelf life before large quarterly batches.",
        "Year 2 mango scale: At 10–20% YoY growth, the April 2028 mango run will need ~28,000–30,000 kg. Begin CM negotiation by December 2027.",
    ]
    for q in questions:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(q)
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc_add_para(doc, "Document version: 1.0  |  Built: May 2026  |  Next review: After Month 2 sales data", italic=True, size=8)

    path = os.path.join(OUT_DIR, "40s-Cookbook-SKU-Demand-Forecast-v1.docx")
    doc.save(path)
    print(f"Word saved  → {path}")
    return path


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

if __name__ == "__main__":
    xl_path   = build_excel()
    docx_path = build_word()
    print("\n✓ Both files exported successfully.")
    print(f"  Excel : {xl_path}")
    print(f"  Word  : {docx_path}")
